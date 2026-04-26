"""
Generate a professional Word document report for the Exam Edge application.
Modeled after the Cybersecurity documentation template structure.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# STYLES SETUP
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.line_spacing = 1.5

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def add_bold_paragraph(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_normal_paragraph(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    doc.add_paragraph()
    return table

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("EXAM EDGE")
run.font.size = Pt(28)
run.bold = True
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("A Comprehensive Competitive Exam Preparation\nWeb Application")
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run("Project Documentation Report")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("April 2026")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (Manual)
# ============================================================
add_heading_custom("TABLE OF CONTENTS", level=1)

toc_entries = [
    ("1.", "Abstract", "3"),
    ("2.", "Introduction", "4"),
    ("3.", "Existing System", "5"),
    ("4.", "Proposed System", "6"),
    ("5.", "System Requirements", "8"),
    ("6.", "System Architecture", "10"),
    ("7.", "Module Description", "12"),
    ("8.", "Database Design", "17"),
    ("9.", "Implementation", "19"),
    ("10.", "Software Environment", "22"),
    ("11.", "Testing", "24"),
    ("12.", "Conclusion", "27"),
    ("13.", "References", "28"),
]

for num, title, page in toc_entries:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}\t{title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(f"\t\t{page}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. ABSTRACT
# ============================================================
add_heading_custom("ABSTRACT", level=1)

add_normal_paragraph(
    "Exam Edge is a full-stack web application designed to provide a comprehensive, "
    "one-stop platform for competitive exam preparation in India. The application addresses "
    "the fragmented nature of exam preparation resources by consolidating exam information, "
    "quiz practice, study materials, current affairs news, and community collaboration into "
    "a single, modern, responsive web interface."
)

add_normal_paragraph(
    "The system is built using a modern technology stack comprising Django REST Framework "
    "for the backend API and React.js with Vite for the frontend. It employs JWT (JSON Web Token) "
    "based authentication for secure user management. The application integrates multiple external "
    "APIs including YouTube Data API for video resources, DuckDuckGo search for PDF study materials, "
    "NewsAPI.ai for live current affairs, and AI-powered quiz generation using Groq (Llama 3.3 70B) "
    "and Google Gemini models."
)

add_normal_paragraph(
    "Key features include an exam catalogue with detailed information about various competitive exams, "
    "an AI-powered practice quiz system that generates questions from uploaded PDF documents, "
    "exam-specific quizzes, community study groups with real-time messaging and material sharing, "
    "a resource aggregation system for videos and PDFs, a live news feed, user dashboard with "
    "personalized notifications, and a responsive design optimized for both desktop and mobile devices."
)

add_normal_paragraph(
    "The platform targets students preparing for examinations such as UPSC, SSC, Banking, Railway, "
    "State PSC, and other government examinations, providing them with an integrated ecosystem for "
    "efficient exam preparation."
)

doc.add_page_break()

# ============================================================
# 2. INTRODUCTION
# ============================================================
add_heading_custom("INTRODUCTION", level=1)

add_normal_paragraph(
    "Competitive examinations form the primary gateway to government employment and higher education "
    "in India. Every year, millions of aspirants prepare for examinations conducted by bodies such as "
    "the Union Public Service Commission (UPSC), Staff Selection Commission (SSC), Institute of Banking "
    "Personnel Selection (IBPS), Railway Recruitment Boards (RRBs), and various State Public Service "
    "Commissions. The sheer volume of candidates, coupled with the diversity and complexity of these "
    "examinations, creates a significant demand for effective preparation tools and resources."
)

add_normal_paragraph(
    "Currently, aspirants rely on a multitude of disconnected platforms for different aspects of their "
    "preparation — separate websites for exam notifications, different apps for quizzes, YouTube for "
    "video lectures, various forums for group discussion, and scattered PDF repositories for study "
    "materials. This fragmented ecosystem leads to inefficiency, information overload, and difficulty "
    "in tracking progress."
)

add_normal_paragraph(
    "Exam Edge was conceived to address this problem by providing a unified, intelligent, and "
    "community-driven platform that brings together all essential aspects of competitive exam "
    "preparation under one roof. The application leverages modern web technologies and artificial "
    "intelligence to deliver a seamless, personalized experience to exam aspirants."
)

add_heading_custom("Objective", level=2)

add_normal_paragraph(
    "The primary objectives of the Exam Edge project are:"
)

objectives = [
    "To develop a centralized web application that aggregates competitive exam information, study resources, practice tools, and community features.",
    "To implement an AI-powered quiz generation system that can create practice questions from uploaded study materials.",
    "To build a community groups system enabling students to collaborate, share materials, and discuss exam strategies.",
    "To integrate real-time news and resource feeds from external APIs for up-to-date exam information.",
    "To provide a secure, scalable, and responsive platform using modern full-stack web development practices.",
    "To implement JWT-based authentication and role-based access control for a secure user experience.",
]
for obj in objectives:
    add_bullet(obj)

doc.add_page_break()

# ============================================================
# 3. EXISTING SYSTEM
# ============================================================
add_heading_custom("EXISTING SYSTEM", level=1)

add_normal_paragraph(
    "The current landscape of competitive exam preparation in India is characterized by several "
    "standalone platforms, each catering to a specific aspect of the preparation journey. Some of "
    "the major existing systems include:"
)

add_heading_custom("Existing Platforms", level=2)

add_bold_paragraph("1. Unacademy / Byju's / Testbook")
add_normal_paragraph(
    "These are major ed-tech platforms offering video courses, live classes, and mock tests. "
    "However, they are primarily commercial, subscription-heavy, and focus on instructor-led "
    "content rather than self-paced, AI-driven learning."
)

add_bold_paragraph("2. Sarkari Result / FreeJobAlert")
add_normal_paragraph(
    "These websites serve as notification aggregators, providing information about exam dates, "
    "results, and admit cards. They lack interactive features such as quizzes, community forums, "
    "or integrated study material."
)

add_bold_paragraph("3. Telegram / WhatsApp Groups")
add_normal_paragraph(
    "Many aspirants rely on social media groups for peer-to-peer learning and material sharing. "
    "However, these platforms lack structure, searchability, and moderation capabilities."
)

add_heading_custom("Disadvantages of the Existing System", level=2)

disadvantages = [
    "Fragmented ecosystem — students must use multiple platforms for different needs.",
    "High cost of subscription on premium platforms.",
    "Lack of AI-powered personalized quiz generation from custom study materials.",
    "No integrated community features within exam preparation platforms.",
    "Information overload without proper categorization and filtering.",
    "Most platforms are consumer-focused and do not provide open, extensible APIs.",
]
for d in disadvantages:
    add_bullet(d)

doc.add_page_break()

# ============================================================
# 4. PROPOSED SYSTEM
# ============================================================
add_heading_custom("PROPOSED SYSTEM", level=1)

add_normal_paragraph(
    "Exam Edge is proposed as a comprehensive, open-source, full-stack web application that "
    "unifies all aspects of competitive exam preparation. The system is designed with a modular "
    "architecture that separates concerns into well-defined Django apps on the backend and "
    "React component-based pages on the frontend."
)

add_heading_custom("Key Features", level=2)

features = [
    ("Exam Catalogue: ", "Browse and search competitive exams across categories (Central, State, Banking, Defence, etc.) with detailed information including eligibility, syllabus, official links, and important dates."),
    ("AI-Powered Practice Quiz: ", "Upload a PDF study material and the system uses Groq Llama 3.3 70B or Google Gemini to generate MCQ questions with explanations. Configurable difficulty (Easy, Medium, Hard) and question count."),
    ("Exam-Specific Quizzes: ", "Pre-built multiple choice question banks for specific exams with scoring, timer, and detailed result analysis."),
    ("Community Groups: ", "Create or join exam-specific study groups with admin approval. Features include group chat/messaging, shared study materials, file attachments, and admin moderation controls."),
    ("Resource Aggregation: ", "YouTube video search integration for exam preparation videos. DuckDuckGo-powered PDF study material search with direct download links."),
    ("Live News Feed: ", "Integration with NewsAPI.ai for real-time exam-related current affairs and notifications."),
    ("User Dashboard: ", "Personalized dashboard showing profile information, target exam, quiz history with scores, and notification center for group activities."),
    ("Notification System: ", "In-app notifications for group join approvals, new messages in groups, and important updates."),
    ("Authentication & Security: ", "JWT-based authentication with access and refresh tokens. Protected routes, CORS configuration, and Django's built-in security middleware."),
]
for bold, text in features:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(bold)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(text)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

add_heading_custom("Advantages of the Proposed System", level=2)

advantages = [
    "All-in-one platform eliminating the need for multiple applications.",
    "AI-powered personalized quiz generation from user-uploaded study materials.",
    "Structured community groups with admin controls, unlike unstructured social media groups.",
    "Free and open-source — no subscription fees.",
    "Modern, responsive UI built with React and TailwindCSS.",
    "Modular, extensible architecture allowing easy addition of new features.",
    "Secure authentication system using industry-standard JWT tokens.",
    "Real-time content from external APIs (YouTube, DuckDuckGo, NewsAPI.ai).",
]
for a in advantages:
    add_bullet(a)

doc.add_page_break()

# ============================================================
# 5. SYSTEM REQUIREMENTS
# ============================================================
add_heading_custom("SYSTEM REQUIREMENTS", level=1)

add_heading_custom("Hardware Requirements", level=2)

add_table(
    ["Component", "Minimum Requirement", "Recommended"],
    [
        ["Processor", "Intel Core i3 / AMD Ryzen 3", "Intel Core i5 / AMD Ryzen 5 or higher"],
        ["RAM", "4 GB", "8 GB or higher"],
        ["Storage", "500 MB free disk space", "1 GB free disk space"],
        ["Network", "Broadband Internet Connection", "High-speed Internet"],
        ["Display", "1366 x 768 resolution", "1920 x 1080 or higher"],
    ]
)

add_heading_custom("Software Requirements", level=2)

add_heading_custom("Backend Dependencies", level=3)
add_table(
    ["Software", "Version", "Purpose"],
    [
        ["Python", "3.10+", "Primary programming language"],
        ["Django", "6.0.3", "Web framework"],
        ["Django REST Framework", "Latest", "RESTful API framework"],
        ["djangorestframework-simplejwt", "Latest", "JWT Authentication"],
        ["django-cors-headers", "Latest", "Cross-Origin Resource Sharing"],
        ["pdfplumber", "Latest", "PDF text extraction"],
        ["Groq Python SDK", "Latest", "AI quiz generation (Llama 3.3 70B)"],
        ["google-genai", "Latest", "AI quiz generation (Gemini fallback)"],
        ["python-dotenv", "Latest", "Environment variable management"],
        ["SQLite3", "Bundled", "Database engine"],
    ]
)

add_heading_custom("Frontend Dependencies", level=3)
add_table(
    ["Software", "Version", "Purpose"],
    [
        ["Node.js", "18+", "JavaScript runtime"],
        ["React", "19.2.4", "UI component library"],
        ["Vite", "8.0.0", "Build tool and dev server"],
        ["React Router DOM", "7.13.1", "Client-side routing"],
        ["Axios", "1.13.6", "HTTP client for API calls"],
        ["jwt-decode", "4.0.0", "JWT token decoding"],
        ["Lucide React", "0.577.0", "Icon library"],
        ["TailwindCSS", "3.4.19", "Utility-first CSS framework"],
    ]
)

add_heading_custom("External API Services", level=3)
add_table(
    ["Service", "Purpose", "Fallback"],
    [
        ["YouTube Data API v3", "Fetch exam preparation videos", "Mock video data"],
        ["DuckDuckGo Search (ddgs)", "Fetch PDF study materials", "Search link fallback"],
        ["NewsAPI.ai (Event Registry)", "Fetch live exam-related news", "Mock news data"],
        ["Groq API (Llama 3.3 70B)", "AI-powered quiz question generation", "Google Gemini fallback"],
        ["Google Gemini API", "Secondary AI quiz generator", "Dummy question fallback"],
    ]
)

doc.add_page_break()

# ============================================================
# 6. SYSTEM ARCHITECTURE
# ============================================================
add_heading_custom("SYSTEM ARCHITECTURE", level=1)

add_normal_paragraph(
    "Exam Edge follows a modern client-server architecture with a clear separation between "
    "the frontend and backend. The system employs a RESTful API design pattern for communication "
    "between the two layers."
)

add_heading_custom("Architecture Overview", level=2)

add_normal_paragraph(
    "The application architecture can be described as a three-tier system:"
)

add_bold_paragraph("1. Presentation Layer (Frontend)")
add_normal_paragraph(
    "Built with React.js and bundled using Vite. The frontend is a Single Page Application (SPA) "
    "that communicates with the backend exclusively through RESTful API calls using Axios. "
    "TailwindCSS is used for responsive, utility-first styling. React Router DOM handles "
    "client-side navigation with protected routes enforcing authentication."
)

add_bold_paragraph("2. Application Layer (Backend)")
add_normal_paragraph(
    "The backend is built using Django 6.0 with Django REST Framework (DRF). It is organized into "
    "modular Django apps: users, exams, quiz, practice_quiz, resources, news, and groups. "
    "Each app encapsulates its own models, views, serializers, and URL configurations. "
    "JWT authentication is implemented using djangorestframework-simplejwt with 3-hour access "
    "tokens and 7-day refresh tokens."
)

add_bold_paragraph("3. Data Layer")
add_normal_paragraph(
    "SQLite3 is used as the database engine for development. Django ORM provides the data access "
    "layer with model definitions for User, Profile, Notification, Exam, Category, ExamDate, "
    "Question, PracticeQuizResult, Group, GroupMember, Message, and Material entities. "
    "Media files (group messages, study materials) are stored in the local filesystem under the "
    "media/ directory."
)

add_heading_custom("API Architecture", level=2)

add_table(
    ["Endpoint Pattern", "App", "Description"],
    [
        ["/api/users/", "Users", "User CRUD, registration, profile, bookmarks, notifications"],
        ["/api/auth/login/", "Auth", "JWT token obtain (login)"],
        ["/api/auth/refresh/", "Auth", "JWT token refresh"],
        ["/api/exams/", "Exams", "Exam catalogue CRUD"],
        ["/api/categories/", "Exams", "Exam categories"],
        ["/api/quiz/", "Quiz", "Exam-specific quiz questions and submission"],
        ["/api/practice-quiz/generate/", "Practice Quiz", "AI-powered quiz generation from PDF"],
        ["/api/practice-quiz/save-result/", "Practice Quiz", "Save quiz attempt results"],
        ["/api/practice-quiz/history/", "Practice Quiz", "User's quiz history"],
        ["/api/resources/videos/", "Resources", "YouTube video search"],
        ["/api/resources/pdfs/", "Resources", "PDF study material search"],
        ["/api/news/", "News", "Live news feed"],
        ["/api/groups/", "Groups", "Group CRUD, join/approve"],
        ["/api/groups/messages/", "Groups", "Group chat messages"],
        ["/api/groups/materials/", "Groups", "Group shared materials"],
    ]
)

add_heading_custom("Frontend Architecture", level=2)

add_table(
    ["Component / Page", "Route", "Description"],
    [
        ["Home", "/", "Landing page with feature overview"],
        ["Login", "/login", "User authentication"],
        ["Register", "/register", "New user registration"],
        ["Dashboard", "/dashboard", "Personalized user dashboard"],
        ["Exams", "/exams", "Exam catalogue with search and filters"],
        ["ExamDetail", "/exams/:id", "Detailed exam information page"],
        ["Quiz", "/quiz", "Exam-specific quiz interface"],
        ["PracticeQuiz", "/practice-quiz", "AI-powered PDF quiz generator"],
        ["Resources", "/resources", "Video and PDF resource aggregator"],
        ["News", "/news", "Live current affairs news feed"],
        ["Groups", "/groups", "Community study groups listing"],
        ["GroupDetail", "/groups/:id", "Group chat, materials, and management"],
    ]
)

doc.add_page_break()

# ============================================================
# 7. MODULE DESCRIPTION
# ============================================================
add_heading_custom("MODULE DESCRIPTION", level=1)

add_normal_paragraph(
    "The Exam Edge application is organized into seven backend modules and a modular frontend. "
    "Each module is designed to encapsulate a specific domain of functionality."
)

# --- Users Module ---
add_heading_custom("Module 1: Users & Authentication", level=2)

add_normal_paragraph(
    "The Users module handles all aspects of user management and authentication. It implements "
    "a custom User model extending Django's AbstractUser, along with a Profile model for "
    "additional user information and a Notification model for in-app notifications."
)

add_bold_paragraph("Key Components:")
components = [
    "Custom User model with AbstractUser extension for future field additions.",
    "Profile model storing state, target exam, and education level.",
    "Notification model for group activity alerts with read/unread status.",
    "UserViewSet providing register, profile, bookmark, and notification endpoints.",
    "JWT authentication using djangorestframework-simplejwt (3-hour access, 7-day refresh).",
    "Permission classes differentiating public (register) and authenticated endpoints.",
]
for c in components:
    add_bullet(c)

# --- Exams Module ---
add_heading_custom("Module 2: Exam Catalogue", level=2)

add_normal_paragraph(
    "The Exam Catalogue module provides a comprehensive database of competitive examinations. "
    "It supports categorization, search, and detailed information for each exam."
)

add_bold_paragraph("Data Model:")
add_normal_paragraph(
    "The module consists of three interconnected models — Category (for grouping exams like "
    "Central, State, Banking, etc.), Exam (storing detailed exam information including name, "
    "conducting body, eligibility, syllabus, official links, application dates, and vacancies), "
    "and ExamDate (enabling multiple date entries per exam such as notification date, application "
    "start/end, exam date, and result date)."
)

# --- Quiz Module ---
add_heading_custom("Module 3: Exam-Specific Quiz", level=2)

add_normal_paragraph(
    "The Quiz module provides pre-built question banks for specific competitive exams. "
    "It features a ReadOnlyModelViewSet that serves 20 randomly ordered questions per request, "
    "with optional filtering by exam. The submit action evaluates user answers against correct "
    "answers and returns a detailed score breakdown with individual question results."
)

# --- Practice Quiz Module ---
add_heading_custom("Module 4: AI-Powered Practice Quiz", level=2)

add_normal_paragraph(
    "The Practice Quiz module is one of the most innovative features of Exam Edge. It allows "
    "users to upload a PDF study material and automatically generates multiple choice questions "
    "using AI models."
)

add_bold_paragraph("Workflow:")
workflow = [
    "User uploads a PDF file and selects difficulty level (Easy, Medium, Hard) and number of questions.",
    "The system extracts text from the PDF using pdfplumber (up to 15,000 characters).",
    "A carefully crafted prompt is sent to Groq's Llama 3.3 70B model for quiz generation.",
    "If Groq fails, the system falls back to Google Gemini 2.5 Flash.",
    "If both AI services fail, fallback dummy questions are returned.",
    "Generated questions include options, correct answer index, and explanations.",
    "Users take the quiz with a timer, and results are saved to the database for history tracking.",
]
for i, w in enumerate(workflow, 1):
    add_bullet(f"Step {i}: {w}")

# --- Resources Module ---
add_heading_custom("Module 5: Resource Aggregation", level=2)

add_normal_paragraph(
    "The Resources module aggregates study materials from two external sources:"
)

add_bold_paragraph("Video Resources (YouTube integration)")
add_normal_paragraph(
    "Uses YouTube Data API v3 to search for exam preparation videos. Returns video title, "
    "thumbnail, channel name, and direct YouTube link. Falls back to mock video data if the "
    "API key is not configured."
)

add_bold_paragraph("PDF Resources (DuckDuckGo integration)")
add_normal_paragraph(
    "Uses the DuckDuckGo search library (ddgs) to find free PDF study materials. The search "
    "query is enhanced with keywords like 'study material', 'notes', 'PDF', and 'free download'. "
    "Returns title, link, and snippet for up to 10 results. Falls back to a DuckDuckGo search "
    "link if the library is unavailable."
)

# --- News Module ---
add_heading_custom("Module 6: Live News Feed", level=2)

add_normal_paragraph(
    "The News module integrates with NewsAPI.ai (Event Registry) to fetch real-time, "
    "English-language articles related to competitive exams. Articles are sorted by date "
    "and include title, description, source, published date, and article URL. A comprehensive "
    "mock dataset serves as fallback when the API is unavailable."
)

# --- Groups Module ---
add_heading_custom("Module 7: Community Groups", level=2)

add_normal_paragraph(
    "The Community Groups module enables collaborative learning through structured study groups. "
    "It is the most feature-rich module in the application."
)

add_bold_paragraph("Group Management:")
group_features = [
    "Group creation with exam type (Central/State) and exam name classification.",
    "Join request system with admin approval workflow.",
    "Custom IsGroupMemberOrAdmin permission class for access control.",
    "Filtering by exam type and exam name.",
]
for g in group_features:
    add_bullet(g)

add_bold_paragraph("Messaging System:")
msg_features = [
    "GroupMember-restricted chat within groups.",
    "File attachment support for messages.",
    "Automatic notification to all other group members on new messages.",
    "Chronological message ordering.",
]
for m in msg_features:
    add_bullet(m)

add_bold_paragraph("Material Sharing:")
mat_features = [
    "Upload and share study materials within a group.",
    "Support for both file uploads and external links.",
    "Member-only access control.",
    "Reverse chronological ordering for latest materials first.",
]
for m in mat_features:
    add_bullet(m)

doc.add_page_break()

# ============================================================
# 8. DATABASE DESIGN
# ============================================================
add_heading_custom("DATABASE DESIGN", level=1)

add_normal_paragraph(
    "The database is implemented using SQLite3 with Django ORM providing the abstraction layer. "
    "The schema consists of 10 primary tables across 4 Django apps."
)

add_heading_custom("Entity-Relationship Summary", level=2)

# Users app tables
add_bold_paragraph("Users App")
add_table(
    ["Table", "Fields", "Relationships"],
    [
        ["User", "id, username, email, password, first_name, last_name, is_active, date_joined", "Extends AbstractUser; 1:1 Profile"],
        ["Profile", "id, user_id, state, target_exam, education_level", "FK → User"],
        ["Notification", "id, user_id, message, is_read, group_id, created_at", "FK → User"],
    ]
)

# Exams app tables
add_bold_paragraph("Exams App")
add_table(
    ["Table", "Fields", "Relationships"],
    [
        ["Category", "id, name, slug", "1:N → Exam"],
        ["Exam", "id, name, category_id, conducting_body, eligibility, syllabus, official_link, apply_link, application_start_date, application_end_date, vacancies, created_at", "FK → Category; 1:N → ExamDate"],
        ["ExamDate", "id, exam_id, date_type, date", "FK → Exam"],
    ]
)

# Quiz app
add_bold_paragraph("Quiz App")
add_table(
    ["Table", "Fields", "Relationships"],
    [
        ["Question", "id, exam_id, text, option_a, option_b, option_c, option_d, correct_answer", "FK → Exam"],
    ]
)

# Practice Quiz app
add_bold_paragraph("Practice Quiz App")
add_table(
    ["Table", "Fields", "Relationships"],
    [
        ["PracticeQuizResult", "id, user_id, title, file_name, score, total_questions, time_taken_seconds, created_at", "FK → User"],
    ]
)

# Groups app
add_bold_paragraph("Groups App")
add_table(
    ["Table", "Fields", "Relationships"],
    [
        ["Group", "id, name, description, exam_type, exam_name, created_at", "1:N → GroupMember, Message, Material"],
        ["GroupMember", "id, group_id, user_id, role, is_approved, joined_at", "FK → Group, User; Unique(group, user)"],
        ["Message", "id, group_id, sender_id, content, file, timestamp", "FK → Group, User"],
        ["Material", "id, group_id, uploaded_by_id, title, file, link, timestamp", "FK → Group, User"],
    ]
)

doc.add_page_break()

# ============================================================
# 9. IMPLEMENTATION
# ============================================================
add_heading_custom("IMPLEMENTATION", level=1)

add_heading_custom("Backend Implementation (Django)", level=2)

add_normal_paragraph(
    "The backend is implemented as a Django project named 'core' with seven registered apps. "
    "The project follows Django best practices with proper separation of concerns — each app "
    "contains its own models.py, views.py, serializers.py, urls.py, and admin.py files."
)

add_bold_paragraph("Project Structure:")
structure_items = [
    "core/ — Project configuration (settings.py, urls.py, wsgi.py, asgi.py)",
    "users/ — Custom user model, profile, notifications, authentication",
    "exams/ — Exam catalogue, categories, exam dates",
    "quiz/ — Pre-built question bank and quiz submission logic",
    "practice_quiz/ — AI-powered PDF quiz generator, results storage",
    "resources/ — YouTube video and PDF study material aggregation",
    "news/ — Live news feed from NewsAPI.ai",
    "groups/ — Community groups, messaging, material sharing",
]
for s in structure_items:
    add_bullet(s)

add_heading_custom("Key Implementation Details", level=3)

add_bold_paragraph("Authentication Flow:")
add_normal_paragraph(
    "Users register via the /api/users/register/ endpoint, which creates both a User and "
    "associated Profile instance. Login is handled by SimpleJWT's TokenObtainPairView, "
    "which returns access and refresh tokens. The frontend stores these tokens in localStorage "
    "and attaches the access token to every API request via an Axios interceptor. Token refresh "
    "is handled automatically when a 401 response is received."
)

add_bold_paragraph("AI Quiz Generation Pipeline:")
add_normal_paragraph(
    "The practice quiz generation follows a multi-fallback strategy. When a user uploads a PDF, "
    "pdfplumber extracts up to 15,000 characters of text. This text is embedded in a structured "
    "prompt requesting JSON-formatted MCQ questions. The system first attempts generation via "
    "Groq's Llama 3.3 70B model, then falls back to Google Gemini 2.5 Flash, and finally to "
    "hardcoded dummy questions. Response parsing handles both raw JSON and markdown-wrapped JSON."
)

add_bold_paragraph("Community Groups Implementation:")
add_normal_paragraph(
    "Groups use a custom IsGroupMemberOrAdmin permission class. When a user creates a group, "
    "they are automatically added as an Admin member with approved status. Join requests create "
    "unapproved GroupMember entries. Admin approval triggers notification creation for the "
    "approved user. Messages generate bulk notifications for all other approved members. "
    "The frontend polls for new messages periodically to simulate real-time chat."
)

add_heading_custom("Frontend Implementation (React)", level=2)

add_normal_paragraph(
    "The frontend is a React 19 Single Page Application bootstrapped with Vite 8. "
    "It uses React Router DOM 7 for client-side routing with a ProtectedRoute component "
    "that checks for valid JWT tokens before rendering authenticated pages."
)

add_bold_paragraph("Component Architecture:")
frontend_items = [
    "App.jsx — Root component with route definitions.",
    "Navbar.jsx — Navigation bar with auth-aware links and notification dropdown.",
    "NotificationDropdown.jsx — Real-time notification display with mark-as-read functionality.",
    "ProtectedRoute.jsx — JWT validation guard for authenticated routes.",
    "axiosConfig.js — Centralized Axios instance with JWT interceptor and auto-refresh logic.",
    "12 page components covering all application features.",
]
for f in frontend_items:
    add_bullet(f)

add_bold_paragraph("State Management:")
add_normal_paragraph(
    "The application uses React's built-in useState and useEffect hooks for local state management. "
    "Authentication state is derived from localStorage tokens. The AuthContext provides "
    "authentication state and methods (login, logout) across the component tree."
)

doc.add_page_break()

# ============================================================
# 10. SOFTWARE ENVIRONMENT
# ============================================================
add_heading_custom("SOFTWARE ENVIRONMENT", level=1)

add_heading_custom("Python", level=2)
add_normal_paragraph(
    "Python is a versatile, high-level programming language known for its clean syntax "
    "and extensive ecosystem. In Exam Edge, Python powers the entire backend through the "
    "Django framework, handling API logic, database operations, PDF parsing, and AI model "
    "integration."
)

add_heading_custom("Django", level=2)
add_normal_paragraph(
    "Django is a high-level Python web framework that encourages rapid development and "
    "clean, pragmatic design. Exam Edge uses Django 6.0 with Django REST Framework to "
    "build RESTful APIs. The framework provides ORM for database operations, built-in "
    "authentication, admin interface, and middleware for security (CSRF, XFrame, etc.)."
)

add_heading_custom("React", level=2)
add_normal_paragraph(
    "React is a JavaScript library for building user interfaces, developed by Meta. "
    "Exam Edge uses React 19 for building a dynamic, component-based frontend. React's "
    "virtual DOM ensures efficient updates, and its hooks API (useState, useEffect) "
    "provides clean state management."
)

add_heading_custom("Vite", level=2)
add_normal_paragraph(
    "Vite is a next-generation frontend build tool that provides lightning-fast hot module "
    "replacement (HMR) and optimized production builds. Exam Edge uses Vite 8 as the build "
    "tool, replacing traditional bundlers like Webpack for significantly faster development "
    "experience."
)

add_heading_custom("TailwindCSS", level=2)
add_normal_paragraph(
    "TailwindCSS is a utility-first CSS framework that enables rapid UI development through "
    "composable utility classes. Exam Edge leverages TailwindCSS 3.4 for responsive design, "
    "consistent styling across components, and a visually appealing interface with custom "
    "colour schemes (india-saffron, india-green, etc.)."
)

add_heading_custom("SQLite", level=2)
add_normal_paragraph(
    "SQLite is a lightweight, serverless, file-based relational database engine. It is used "
    "in Exam Edge for development and prototyping, providing zero-configuration database "
    "setup. The schema supports all application requirements through Django ORM migrations."
)

add_heading_custom("Development Tools", level=2)
add_table(
    ["Tool", "Purpose"],
    [
        ["VS Code", "Primary IDE for development"],
        ["Git & GitHub", "Version control and remote repository"],
        ["Postman", "API testing"],
        ["Chrome DevTools", "Frontend debugging"],
        ["Python venv", "Virtual environment for backend dependencies"],
        ["npm", "Frontend package management"],
    ]
)

doc.add_page_break()

# ============================================================
# 11. TESTING
# ============================================================
add_heading_custom("TESTING", level=1)

add_normal_paragraph(
    "The purpose of testing is to discover errors. Testing is the process of trying to "
    "discover every conceivable fault or weakness in a work product. It provides a way to "
    "check the functionality of components, sub-assemblies, assemblies, and the finished "
    "product. The testing methodology applied to Exam Edge encompasses multiple levels of "
    "verification."
)

add_heading_custom("Testing Objective", level=2)

add_normal_paragraph(
    "Software testing is an important phase in the development of the system. The objective "
    "while testing the application is to ensure that the application runs error-free by "
    "verifying all modules and their integrations. The following types of testing were performed:"
)

testing_types = [
    "Unit Testing",
    "Integration Testing",
    "User Interface Testing",
    "Functional Testing",
    "System Testing",
    "API Testing",
]
for t in testing_types:
    add_bullet(t)

add_heading_custom("Unit Testing", level=2)
add_normal_paragraph(
    "This is the first level of testing, where individual modules are tested against their "
    "specifications. Each Django app (users, exams, quiz, practice_quiz, resources, news, "
    "groups) was tested independently to verify correct model definitions, serializer "
    "validation, and view logic. Individual React components were tested to ensure proper "
    "rendering and state management."
)

add_heading_custom("Integration Testing", level=2)
add_normal_paragraph(
    "Integration testing verifies the interaction between modules. Key integration points "
    "tested include: frontend-to-backend API communication via Axios, JWT authentication "
    "flow from login to protected API access, AI quiz generation pipeline from PDF upload "
    "through text extraction to AI model response, notification creation when group actions "
    "occur (join approval, new messages), and resource API fallback mechanisms from live "
    "APIs to mock data."
)

add_heading_custom("User Interface Testing", level=2)
add_normal_paragraph(
    "UI testing was performed across multiple browsers (Chrome, Firefox, Edge) and screen "
    "sizes to ensure responsive design. Tests verified proper navigation, form validation, "
    "loading states, error handling, and visual consistency. Special attention was given to "
    "the quiz interface (timer, question navigation, answer selection) and group chat "
    "interface (message display, file uploads)."
)

add_heading_custom("API Testing", level=2)
add_normal_paragraph(
    "All REST API endpoints were tested using Python's requests library and Postman. "
    "Test scripts (test_api.py) verify correct HTTP status codes, response formats, "
    "authentication requirements, and error handling for edge cases. Tests cover both "
    "happy paths and error scenarios such as invalid tokens, missing required fields, "
    "and unauthorized access attempts."
)

add_heading_custom("Performance Testing", level=2)
add_normal_paragraph(
    "Performance testing verified the application's responsiveness under normal usage "
    "conditions. Key metrics include API response times (all endpoints respond within "
    "500ms for database queries), PDF processing time for AI quiz generation (varies "
    "based on document size and AI model load), and frontend page load times with lazy "
    "loading and efficient re-rendering via React's virtual DOM."
)

add_heading_custom("System Testing", level=2)
add_normal_paragraph(
    "After all modules were integrated, comprehensive system testing was performed to "
    "verify the complete user journey: registration → login → browse exams → take quiz → "
    "upload PDF for practice quiz → join groups → send messages → share materials → "
    "receive notifications → logout. System testing confirmed that all features work "
    "harmoniously and data flows correctly across all modules."
)

doc.add_page_break()

# ============================================================
# 12. CONCLUSION
# ============================================================
add_heading_custom("CONCLUSION", level=1)

add_normal_paragraph(
    "Exam Edge successfully demonstrates the development of a comprehensive, full-stack "
    "web application for competitive exam preparation. The project integrates modern web "
    "technologies (Django, React, Vite) with artificial intelligence (Groq Llama 3.3 70B, "
    "Google Gemini) and external data sources (YouTube, DuckDuckGo, NewsAPI.ai) to deliver "
    "a feature-rich platform for exam aspirants."
)

add_normal_paragraph(
    "The application addresses the core problem of fragmented exam preparation resources by "
    "providing a unified platform that combines exam information, AI-powered quiz generation, "
    "study material aggregation, live news, and community collaboration. The modular architecture "
    "ensures maintainability and extensibility, while JWT-based authentication provides secure "
    "access control."
)

add_normal_paragraph(
    "Key technical achievements of this project include:"
)

achievements = [
    "Implementation of a multi-fallback AI quiz generation pipeline (Groq → Gemini → Fallback).",
    "Integration of three external APIs with graceful fallback mechanisms.",
    "A complete community groups system with admin approval workflow and in-app notifications.",
    "A responsive, modern UI using React 19 and TailwindCSS with custom Indian-themed colour palette.",
    "Secure REST API architecture with JWT authentication and role-based permissions.",
    "Clean, modular codebase with 7 backend apps and 12+ frontend pages.",
]
for a in achievements:
    add_bullet(a)

add_heading_custom("Future Enhancements", level=2)

future = [
    "WebSocket integration for real-time group chat instead of polling.",
    "PostgreSQL migration for production-grade database performance.",
    "User progress tracking and analytics dashboard with charts.",
    "Spaced repetition algorithm for intelligent quiz scheduling.",
    "Mobile application using React Native with shared API backend.",
    "OAuth integration for Google and social media login.",
    "Admin dashboard for content management and user moderation.",
    "PDF annotation and note-taking within the application.",
]
for f in future:
    add_bullet(f)

doc.add_page_break()

# ============================================================
# 13. REFERENCES
# ============================================================
add_heading_custom("REFERENCES", level=1)

references = [
    "Django Software Foundation. \"Django Documentation - Version 6.0.\" https://docs.djangoproject.com/en/6.0/",
    "Django REST Framework. \"Django REST Framework Documentation.\" https://www.django-rest-framework.org/",
    "Meta Platforms, Inc. \"React Documentation - Version 19.\" https://react.dev/",
    "Evan You et al. \"Vite - Next Generation Frontend Tooling.\" https://vitejs.dev/",
    "Adam Wathan et al. \"TailwindCSS Documentation - Version 3.\" https://tailwindcss.com/docs",
    "SimpleJWT. \"Django REST Framework Simple JWT.\" https://django-rest-framework-simplejwt.readthedocs.io/",
    "Groq, Inc. \"Groq API Documentation - Llama 3.3 70B.\" https://groq.com/",
    "Google. \"Google Gemini API Documentation.\" https://ai.google.dev/",
    "YouTube. \"YouTube Data API v3 Reference.\" https://developers.google.com/youtube/v3",
    "DuckDuckGo. \"DuckDuckGo Search API.\" https://duckduckgo.com/",
    "NewsAPI.ai. \"Event Registry API Documentation.\" https://eventregistry.org/",
    "SQLite Consortium. \"SQLite Documentation.\" https://www.sqlite.org/docs.html",
    "Axios. \"Axios HTTP Client Documentation.\" https://axios-http.com/docs/intro",
    "React Router. \"React Router DOM Documentation - Version 7.\" https://reactrouter.com/",
    "pdfplumber. \"pdfplumber - PDF Parsing for Python.\" https://github.com/jsvine/pdfplumber",
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {ref}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ============================================================
# SAVE DOCUMENT
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Documentation - Exam Edge.docx")
doc.save(output_path)
print(f"Document saved successfully to: {output_path}")
